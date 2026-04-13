import re
from io import BytesIO

import pandas as pd
import plotly.express as px
import streamlit as st
from docx import Document
from fpdf import FPDF

st.set_page_config(page_title="Audit Dashboard", layout="wide")


def load_doc(file):
    file_bytes = file.getvalue()
    return Document(BytesIO(file_bytes))


def extract_summary_table(file):
    doc = load_doc(file)
    data = []

    for table in doc.tables:
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]

            if len(row_data) >= 5 and ("Failed" in row_data or "Passed" in row_data):
                try:
                    data.append(
                        {
                            "Asset": row_data[1],
                            "Vulnerability": row_data[2],
                            "Status": row_data[3],
                            "Recommendation": row_data[4],
                            "Reference": row_data[5] if len(row_data) > 5 else "",
                        }
                    )
                except Exception:
                    continue

    return pd.DataFrame(data)


def extract_summary_counts(file):
    doc = load_doc(file)
    text = "\n".join([p.text for p in doc.paragraphs])

    passed, failed = 0, 0

    passed_match = re.search(r"(\d+)\s*Passed", text)
    failed_match = re.search(r"(\d+)\s*Failed", text)

    if passed_match:
        passed = int(passed_match.group(1))
    if failed_match:
        failed = int(failed_match.group(1))

    return passed, failed


def generate_ai_recommendation(vuln):
    vuln = vuln.lower()

    if "ssh" in vuln:
        return "Enable SSH and disable Telnet."
    if "vlan" in vuln:
        return "Ensure VLAN segmentation."
    if "password" in vuln:
        return "Use SHA-512 encryption."
    if "telnet" in vuln:
        return "Disable Telnet."
    return "Follow CIS/NIST standards."


def generate_pdf(summary_text):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.multi_cell(0, 10, summary_text)
    return pdf.output(dest="S").encode("latin-1")


st.title("Audit Dashboard")

uploaded_files = st.file_uploader(
    "Upload Audit Report (.docx)", type=["docx"], accept_multiple_files=True
)

if not uploaded_files:
    st.info("Upload one or more .docx audit reports to view the dashboard.")
else:
    all_data = []
    total_passed = 0
    total_failed = 0

    for file in uploaded_files:
        df = extract_summary_table(file)
        passed, failed = extract_summary_counts(file)

        total_passed += passed
        total_failed += failed

        if not df.empty:
            df["Source File"] = file.name
            all_data.append(df)

    if not all_data:
        st.warning("No findings were detected in the uploaded reports.")
    else:
        df = pd.concat(all_data, ignore_index=True)

        df["Compliance"] = df["Status"].apply(
            lambda x: "Non-Compliant" if "Fail" in str(x) else "Compliant"
        )

        def risk_score(status):
            status = str(status)
            if "Fail" in status:
                return 9
            if "Partial" in status:
                return 5
            return 1

        df["Risk Score"] = df["Status"].apply(risk_score)

        def risk_level(score):
            if score >= 7:
                return "Critical"
            if score >= 4:
                return "Medium"
            return "Low"

        df["Risk Level"] = df["Risk Score"].apply(risk_level)
        df["AI Recommendation"] = df["Vulnerability"].apply(generate_ai_recommendation)

        total = total_passed + total_failed
        compliance = round((total_passed / total) * 100, 2) if total > 0 else 0

        col1, col2, col3 = st.columns(3)
        col1.metric("Passed", total_passed)
        col2.metric("Failed", total_failed)
        col3.metric("Compliance %", compliance)

        st.divider()

        tab1, tab2, tab3 = st.tabs(["Charts", "Findings", "Risk"])

        with tab1:
            fig1 = px.pie(df, names="Compliance", title="Compliance Split")
            st.plotly_chart(fig1, use_container_width=True)

        with tab2:
            st.dataframe(df, use_container_width=True)

        with tab3:
            fig3 = px.bar(
                df,
                x="Vulnerability",
                y="Risk Score",
                color="Risk Level",
                title="Risk Scores by Vulnerability",
            )
            st.plotly_chart(fig3, use_container_width=True)

        summary = f"""
Total Findings: {len(df)}
Passed: {total_passed}
Failed: {total_failed}
Compliance: {compliance}%
"""

        st.text_area("Executive Summary", summary, height=140)

        pdf_bytes = generate_pdf(summary)
        st.download_button(
            "Download PDF",
            data=pdf_bytes,
            file_name="audit_report.pdf",
            mime="application/pdf",
        )
